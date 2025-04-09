#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Report generator module for the SEO Analysis Tool.
Contains classes for generating reports.
"""

import logging
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import datetime
import io
from wordcloud import WordCloud
import matplotlib.cm as cm
import matplotlib.colors as mcolors

class ReportGenerator:
    """Class for generating reports."""
    
    def __init__(self, output_dir="reports"):
        """
        Initialize the report generator.
        
        Args:
            output_dir (str): The output directory for reports
        """
        self.output_dir = output_dir
        self.logger = logging.getLogger(__name__)
        
        # Create the output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)
    
    def generate_seo_report(self, url, meta_analyzer, image_analyzer, links_analyzer, reachability_analyzer, topic_analyzer, report_title=None):
        """
        Generate an SEO report.
        
        Args:
            url (str): The URL
            meta_analyzer (MetaDescriptionAnalyzer): The meta description analyzer
            image_analyzer (ImageAnalyzer): The image analyzer
            links_analyzer (LinksAnalyzer): The links analyzer
            reachability_analyzer (ReachabilityAnalyzer): The reachability analyzer
            topic_analyzer (MainTopicAnalyzer): The main topic analyzer
            report_title (str, optional): Custom report title. If None, a default title will be used.
        
        Returns:
            str: The path to the generated report
        """
        self.logger.info(f"Generating SEO report for URL: {url}")
        print(f"Starting SEO report generation for URL: {url}...")
        
        # Create a new document
        doc = Document()
        print("Report document created (10% complete)")
        
        # Add styles
        self._add_styles(doc)
        print("Styles added (15% complete)")
        
        # Add title
        if report_title:
            self._add_title(doc, report_title)
        else:
            self._add_title(doc, f"SEO Analysis Report: {url}")
        
        # Add date
        self._add_date(doc)
        print("Title and date added (20% complete)")
        
        # Add summary
        print("Generating summary...")
        self._add_summary(doc, url, meta_analyzer, image_analyzer, links_analyzer, reachability_analyzer, topic_analyzer)
        print("Summary added (30% complete)")
        
        # Add meta description analysis
        print("Processing meta description analysis...")
        meta_desc_df = meta_analyzer.to_dataframe()
        self._add_meta_description_analysis(doc, meta_analyzer)
        print(f"Added meta description analysis for {len(meta_desc_df)} URLs (40% complete)")
        
        # Add image analysis
        print("Processing image analysis...")
        image_df = image_analyzer.to_dataframe()
        self._add_image_analysis(doc, image_analyzer)
        print(f"Added image analysis for {len(image_df)} images (50% complete)")
        
        # Add links analysis
        print("Processing links analysis...")
        links_df = links_analyzer.to_dataframe()
        self._add_links_analysis(doc, links_analyzer)
        print(f"Added links analysis for {len(links_df)} links (60% complete)")
        
        # Add reachability analysis
        print("Processing reachability analysis...")
        reachability_df = reachability_analyzer.to_dataframe()
        self._add_reachability_analysis(doc, reachability_analyzer)
        print(f"Added reachability analysis for {len(reachability_df)} URLs (70% complete)")
        
        # Add topic analysis
        print("Processing topic analysis...")
        topic_df = topic_analyzer.to_dataframe()
        self._add_topic_analysis(doc, topic_analyzer)
        print(f"Added topic analysis for {len(topic_df)} URLs (80% complete)")
        
        # Add recommendations
        print("Generating recommendations...")
        self._add_recommendations(doc, meta_analyzer, image_analyzer, links_analyzer, reachability_analyzer, topic_analyzer)
        print("Recommendations added (90% complete)")
        
        # Save the document
        print("Saving report document...")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"seo_report_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)
        
        doc.save(filepath)
        
        self.logger.info(f"SEO report saved to: {filepath}")
        print(f"SEO report completed (100%)")
        print(f"Report saved to: {filepath}")
        
        return filepath
    
    def generate_search_console_report(self, search_console_analyzer):
        """
        Generate a Search Console report.
        
        Args:
            search_console_analyzer (SearchConsoleAnalyzer): The Search Console analyzer
        
        Returns:
            str: The path to the generated report
        """
        self.logger.info("Generating Search Console report")
        print("Starting Search Console report generation...")
        
        # Create a new document
        doc = Document()
        print("Report document created (10% complete)")
        
        # Add styles
        self._add_styles(doc)
        print("Styles added (20% complete)")
        
        # Add title
        self._add_title(doc, "Search Console Analysis Report")
        
        # Add date
        self._add_date(doc)
        print("Title and date added (30% complete)")
        
        # Add clusters analysis
        print("Processing query clusters...")
        clusters = search_console_analyzer.get_clusters()
        self._add_clusters_analysis(doc, clusters, "Query")
        print(f"Added {len(clusters)} query clusters to report (50% complete)")
        
        # Add topics analysis
        print("Processing URL topics...")
        topics = search_console_analyzer.get_topics()
        self._add_topics_analysis(doc, topics)
        print(f"Added topics for {len(topics)} URLs to report (70% complete)")
        
        # Add internal link suggestions
        print("Generating internal link suggestions...")
        suggestions = search_console_analyzer.suggest_internal_links()
        self._add_internal_link_suggestions(doc, suggestions)
        print(f"Added {len(suggestions)} internal link suggestions to report (90% complete)")
        
        # Save the document
        print("Saving report document...")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"search_console_report_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)
        
        doc.save(filepath)
        
        self.logger.info(f"Search Console report saved to: {filepath}")
        print(f"Search Console report completed (100%)")
        print(f"Report saved to: {filepath}")
        
        return filepath
    
    def generate_semrush_report(self, semrush_analyzer):
        """
        Generate a SEMrush report.
        
        Args:
            semrush_analyzer (SEMrushAnalyzer): The SEMrush analyzer
        
        Returns:
            str: The path to the generated report
        """
        self.logger.info("Generating SEMrush report")
        print("Starting SEMrush report generation...")
        
        # Create a new document
        doc = Document()
        print("Report document created (10% complete)")
        
        # Add styles
        self._add_styles(doc)
        print("Styles added (20% complete)")
        
        # Add title
        self._add_title(doc, "SEMrush Analysis Report")
        
        # Add date
        self._add_date(doc)
        print("Title and date added (30% complete)")
        
        # Add clusters analysis
        print("Processing keyword clusters...")
        clusters = semrush_analyzer.get_clusters()
        self._add_clusters_analysis(doc, clusters, "Keyword")
        print(f"Added {len(clusters)} keyword clusters to report (45% complete)")
        
        # Add topics analysis
        print("Processing URL topics...")
        topics = semrush_analyzer.get_topics()
        self._add_topics_analysis(doc, topics)
        print(f"Added topics for {len(topics)} URLs to report (60% complete)")
        
        # Add visibility and traffic analysis
        print("Analyzing visibility and traffic...")
        visibility = semrush_analyzer.get_visibility()
        traffic = semrush_analyzer.get_traffic()
        self._add_visibility_traffic_analysis(doc, visibility, traffic)
        print("Added visibility and traffic analysis to report (75% complete)")
        
        # Add internal link suggestions
        print("Generating internal link suggestions...")
        suggestions = semrush_analyzer.suggest_internal_links()
        self._add_internal_link_suggestions(doc, suggestions)
        print(f"Added {len(suggestions)} internal link suggestions to report (90% complete)")
        
        # Save the document
        print("Saving report document...")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"semrush_report_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)
        
        doc.save(filepath)
        
        self.logger.info(f"SEMrush report saved to: {filepath}")
        print(f"SEMrush report completed (100%)")
        print(f"Report saved to: {filepath}")
        
        return filepath
    
    def generate_comparison_report(self, search_console_comparison):
        """
        Generate a comparison report.
        
        Args:
            search_console_comparison (SearchConsoleComparison): The Search Console comparison
        
        Returns:
            str: The path to the generated report
        """
        self.logger.info("Generating comparison report")
        print("Starting comparison report generation...")
        
        # Create a new document
        doc = Document()
        print("Report document created (10% complete)")
        
        # Add styles
        self._add_styles(doc)
        print("Styles added (20% complete)")
        
        # Add title
        self._add_title(doc, "Search Console Comparison Report")
        
        # Add date
        self._add_date(doc)
        print("Title and date added (30% complete)")
        
        # Add query comparison
        print("Processing query comparison...")
        query_comparison = search_console_comparison.get_query_comparison()
        self._add_query_comparison(doc, query_comparison)
        print(f"Added query comparison with {len(query_comparison) if query_comparison is not None else 0} queries (40% complete)")
        
        # Add landing page comparison
        print("Processing landing page comparison...")
        landing_page_comparison = search_console_comparison.get_landing_page_comparison()
        self._add_landing_page_comparison(doc, landing_page_comparison)
        print(f"Added landing page comparison with {len(landing_page_comparison) if landing_page_comparison is not None else 0} pages (50% complete)")
        
        # Add improved queries
        print("Processing improved queries...")
        improved_queries = search_console_comparison.get_improved_queries()
        self._add_improved_queries(doc, improved_queries)
        print(f"Added {len(improved_queries) if improved_queries is not None else 0} improved queries (60% complete)")
        
        # Add declined queries
        print("Processing declined queries...")
        declined_queries = search_console_comparison.get_declined_queries()
        self._add_declined_queries(doc, declined_queries)
        print(f"Added {len(declined_queries) if declined_queries is not None else 0} declined queries (70% complete)")
        
        # Add improved landing pages
        print("Processing improved landing pages...")
        improved_landing_pages = search_console_comparison.get_improved_landing_pages()
        self._add_improved_landing_pages(doc, improved_landing_pages)
        print(f"Added {len(improved_landing_pages) if improved_landing_pages is not None else 0} improved landing pages (80% complete)")
        
        # Add declined landing pages
        print("Processing declined landing pages...")
        declined_landing_pages = search_console_comparison.get_declined_landing_pages()
        self._add_declined_landing_pages(doc, declined_landing_pages)
        print(f"Added {len(declined_landing_pages) if declined_landing_pages is not None else 0} declined landing pages (90% complete)")
        
        # Save the document
        print("Saving report document...")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"comparison_report_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)
        
        doc.save(filepath)
        
        self.logger.info(f"Comparison report saved to: {filepath}")
        print(f"Comparison report completed (100%)")
        print(f"Report saved to: {filepath}")
        
        return filepath
    
    def generate_final_report(self, url, meta_analyzer, image_analyzer, links_analyzer, reachability_analyzer, topic_analyzer, search_console_analyzer=None, semrush_analyzer=None):
        """
        Generate a final report.
        
        Args:
            url (str): The URL
            meta_analyzer (MetaDescriptionAnalyzer): The meta description analyzer
            image_analyzer (ImageAnalyzer): The image analyzer
            links_analyzer (LinksAnalyzer): The links analyzer
            reachability_analyzer (ReachabilityAnalyzer): The reachability analyzer
            topic_analyzer (MainTopicAnalyzer): The main topic analyzer
            search_console_analyzer (SearchConsoleAnalyzer, optional): The Search Console analyzer
            semrush_analyzer (SEMrushAnalyzer, optional): The SEMrush analyzer
        
        Returns:
            str: The path to the generated report
        """
        self.logger.info(f"Generating final report for URL: {url}")
        print(f"Starting final report generation for URL: {url}...")
        
        # Create a new document
        doc = Document()
        print("Report document created (5% complete)")
        
        # Add styles
        self._add_styles(doc)
        print("Styles added (10% complete)")
        
        # Add title
        self._add_title(doc, f"SEO Final Report: {url}")
        
        # Add date
        self._add_date(doc)
        print("Title and date added (15% complete)")
        
        # Add executive summary
        print("Generating executive summary...")
        self._add_executive_summary(doc, url, meta_analyzer, image_analyzer, links_analyzer, reachability_analyzer, topic_analyzer, search_console_analyzer, semrush_analyzer)
        print("Executive summary added (20% complete)")
        
        # Add SEO analysis
        doc.add_heading("SEO Analysis", level=1)
        
        # Add meta description analysis
        print("Processing meta description analysis...")
        self._add_meta_description_analysis(doc, meta_analyzer)
        print("Meta description analysis added (30% complete)")
        
        # Add image analysis
        print("Processing image analysis...")
        self._add_image_analysis(doc, image_analyzer)
        print("Image analysis added (40% complete)")
        
        # Add links analysis
        print("Processing links analysis...")
        self._add_links_analysis(doc, links_analyzer)
        print("Links analysis added (50% complete)")
        
        # Add reachability analysis
        print("Processing reachability analysis...")
        self._add_reachability_analysis(doc, reachability_analyzer)
        print("Reachability analysis added (60% complete)")
        
        # Add topic analysis
        print("Processing topic analysis...")
        self._add_topic_analysis(doc, topic_analyzer)
        print("Topic analysis added (70% complete)")
        
        # Add Search Console analysis if available
        if search_console_analyzer:
            print("Processing Search Console analysis...")
            doc.add_heading("Search Console Analysis", level=1)
            self._add_clusters_analysis(doc, search_console_analyzer.get_clusters(), "Query")
            self._add_topics_analysis(doc, search_console_analyzer.get_topics())
            print("Search Console analysis added (75% complete)")
        
        # Add SEMrush analysis if available
        if semrush_analyzer:
            print("Processing SEMrush analysis...")
            doc.add_heading("SEMrush Analysis", level=1)
            self._add_clusters_analysis(doc, semrush_analyzer.get_clusters(), "Keyword")
            self._add_topics_analysis(doc, semrush_analyzer.get_topics())
            self._add_visibility_traffic_analysis(doc, semrush_analyzer.get_visibility(), semrush_analyzer.get_traffic())
            print("SEMrush analysis added (80% complete)")
        
        # Add internal link suggestions
        print("Generating internal link suggestions...")
        doc.add_heading("Internal Link Suggestions", level=1)
        
        if search_console_analyzer:
            suggestions = search_console_analyzer.suggest_internal_links()
            self._add_internal_link_suggestions(doc, suggestions, "Search Console")
            print(f"Added {len(suggestions)} Search Console internal link suggestions")
        
        if semrush_analyzer:
            suggestions = semrush_analyzer.suggest_internal_links()
            self._add_internal_link_suggestions(doc, suggestions, "SEMrush")
            print(f"Added {len(suggestions)} SEMrush internal link suggestions")
        
        print("Internal link suggestions added (90% complete)")
        
        # Add recommendations
        print("Generating recommendations...")
        self._add_recommendations(doc, meta_analyzer, image_analyzer, links_analyzer, reachability_analyzer, topic_analyzer, search_console_analyzer, semrush_analyzer)
        print("Recommendations added (95% complete)")
        
        # Save the document
        print("Saving report document...")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"final_report_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)
        
        doc.save(filepath)
        
        self.logger.info(f"Final report saved to: {filepath}")
        print(f"Final report completed (100%)")
        print(f"Report saved to: {filepath}")
        
        return filepath
    
    def _add_styles(self, doc):
        """
        Add styles to the document.
        
        Args:
            doc (Document): The document
        """
        # Add title style if it doesn't exist
        if "Title" not in doc.styles:
            title_style = doc.styles.add_style("Title", WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = "Arial"
            title_style.font.size = Pt(24)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(0, 0, 128)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
        else:
            # Use existing title style
            title_style = doc.styles["Title"]
        
        # Add heading 1 style
        heading1_style = doc.styles["Heading 1"]
        heading1_style.font.name = "Arial"
        heading1_style.font.size = Pt(18)
        heading1_style.font.bold = True
        heading1_style.font.color.rgb = RGBColor(0, 0, 128)
        heading1_style.paragraph_format.space_before = Pt(12)
        heading1_style.paragraph_format.space_after = Pt(6)
        
        # Add heading 2 style
        heading2_style = doc.styles["Heading 2"]
        heading2_style.font.name = "Arial"
        heading2_style.font.size = Pt(14)
        heading2_style.font.bold = True
        heading2_style.font.color.rgb = RGBColor(0, 0, 128)
        heading2_style.paragraph_format.space_before = Pt(12)
        heading2_style.paragraph_format.space_after = Pt(6)
        
        # Add normal style
        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Arial"
        normal_style.font.size = Pt(11)
        normal_style.paragraph_format.space_after = Pt(6)
    
    def _add_title(self, doc, title):
        """
        Add a title to the document.
        
        Args:
            doc (Document): The document
            title (str): The title
        """
        doc.add_paragraph(title, style="Title")
    
    def _add_date(self, doc):
        """
        Add the current date to the document.
        
        Args:
            doc (Document): The document
        """
        date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        date_paragraph = doc.add_paragraph(f"Generated on: {date}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    
    def _add_summary(self, doc, url, meta_analyzer, image_analyzer, links_analyzer, reachability_analyzer, topic_analyzer):
        """
        Add a summary to the document.
        
        Args:
            doc (Document): The document
            url (str): The URL
            meta_analyzer (MetaDescriptionAnalyzer): The meta description analyzer
            image_analyzer (ImageAnalyzer): The image analyzer
            links_analyzer (LinksAnalyzer): The links analyzer
            reachability_analyzer (ReachabilityAnalyzer): The reachability analyzer
            topic_analyzer (MainTopicAnalyzer): The main topic analyzer
        """
        doc.add_heading("Summary", level=1)
        
        # Add URL
        doc.add_paragraph(f"URL: {url}")
        
        # Add meta description summary
        meta_desc_df = meta_analyzer.to_dataframe()
        urls_with_meta_desc = len(meta_desc_df[meta_desc_df["Has Meta Description"]])
        urls_without_meta_desc = len(meta_desc_df) - urls_with_meta_desc
        
        doc.add_paragraph(f"Meta Descriptions: {urls_with_meta_desc} URLs have meta descriptions, {urls_without_meta_desc} URLs don't have meta descriptions.")
        
        # Add image summary
        image_df = image_analyzer.to_dataframe()
        images_with_alt = len(image_df[image_df["Has Alt Text"]])
        images_without_alt = len(image_df) - images_with_alt
        
        doc.add_paragraph(f"Images: {len(image_df)} images found, {images_with_alt} have alt text, {images_without_alt} don't have alt text.")
        
        # Add links summary
        links_df = links_analyzer.to_dataframe()
        internal_links = len(links_df[links_df["Is Internal"]])
        external_links = len(links_df[links_df["Is External"]])
        broken_links = len(links_df[links_df["Is Broken"]])
        
        doc.add_paragraph(f"Links: {len(links_df)} links found, {internal_links} internal links, {external_links} external links, {broken_links} broken links.")
        
        # Add reachability summary
        reachability_df = reachability_analyzer.to_dataframe()
        reachable_urls = len(reachability_df[reachability_df["Is Reachable"]])
        orphan_pages = len(reachability_df[reachability_df["Is Orphan Page"]])
        
        doc.add_paragraph(f"Reachability: {reachable_urls} URLs are reachable from the home page, {orphan_pages} URLs are orphan pages.")
        
        # Add topic summary
        topic_df = topic_analyzer.to_dataframe()
        
        doc.add_paragraph(f"Topics: {len(topic_df)} URLs analyzed for topics.")
    
    def _add_meta_description_analysis(self, doc, meta_analyzer):
        """
        Add meta description analysis to the document.
        
        Args:
            doc (Document): The document
            meta_analyzer (MetaDescriptionAnalyzer): The meta description analyzer
        """
        doc.add_heading("Meta Description Analysis", level=1)
        
        # Get the data
        meta_desc_df = meta_analyzer.to_dataframe()
        
        # Add summary
        urls_with_meta_desc = len(meta_desc_df[meta_desc_df["Has Meta Description"]])
        urls_without_meta_desc = len(meta_desc_df) - urls_with_meta_desc
        
        doc.add_paragraph(f"URLs with meta descriptions: {urls_with_meta_desc}")
        doc.add_paragraph(f"URLs without meta descriptions: {urls_without_meta_desc}")
        
        # Add URLs without meta descriptions
        if urls_without_meta_desc > 0:
            doc.add_heading("URLs Without Meta Descriptions", level=2)
            
            urls_without_meta_desc_df = meta_desc_df[~meta_desc_df["Has Meta Description"]]
            
            for _, row in urls_without_meta_desc_df.iterrows():
                doc.add_paragraph(row["URL"], style="List Bullet")
        
        # Add URLs with short meta descriptions
        urls_with_short_meta_desc = len(meta_desc_df[meta_desc_df["Meta Description Quality"] == "too_short"])
        
        if urls_with_short_meta_desc > 0:
            doc.add_heading("URLs With Short Meta Descriptions", level=2)
            
            urls_with_short_meta_desc_df = meta_desc_df[meta_desc_df["Meta Description Quality"] == "too_short"]
            
            for _, row in urls_with_short_meta_desc_df.iterrows():
                doc.add_paragraph(f"{row['URL']} ({row['Meta Description Length']} characters)", style="List Bullet")
        
        # Add URLs with long meta descriptions
        urls_with_long_meta_desc = len(meta_desc_df[meta_desc_df["Meta Description Quality"] == "too_long"])
        
        if urls_with_long_meta_desc > 0:
            doc.add_heading("URLs With Long Meta Descriptions", level=2)
            
            urls_with_long_meta_desc_df = meta_desc_df[meta_desc_df["Meta Description Quality"] == "too_long"]
            
            for _, row in urls_with_long_meta_desc_df.iterrows():
                doc.add_paragraph(f"{row['URL']} ({row['Meta Description Length']} characters)", style="List Bullet")
    
    def _add_image_analysis(self, doc, image_analyzer):
        """
        Add image analysis to the document.
        
        Args:
            doc (Document): The document
            image_analyzer (ImageAnalyzer): The image analyzer
        """
        doc.add_heading("Image Analysis", level=1)
        
        # Get the data
        image_df = image_analyzer.to_dataframe()
        
        # Add summary
        images_with_alt = len(image_df[image_df["Has Alt Text"]])
        images_without_alt = len(image_df) - images_with_alt
        
        doc.add_paragraph(f"Images with alt text: {images_with_alt}")
        doc.add_paragraph(f"Images without alt text: {images_without_alt}")
        
        # Add images without alt text
        if images_without_alt > 0:
            doc.add_heading("Images Without Alt Text", level=2)
            
            images_without_alt_df = image_df[~image_df["Has Alt Text"]]
            
            for _, row in images_without_alt_df.iterrows():
                doc.add_paragraph(f"{row['URL']}: {row['Image Source']}", style="List Bullet")
        
        # Add large images
        large_images = len(image_df[image_df["Image Size"] > 100000])  # > 100 KB
        
        if large_images > 0:
            doc.add_heading("Large Images", level=2)
            
            large_images_df = image_df[image_df["Image Size"] > 100000]
            
            for _, row in large_images_df.iterrows():
                doc.add_paragraph(f"{row['URL']}: {row['Image Source']} ({row['Image Size'] / 1024:.1f} KB)", style="List Bullet")
    
    def _add_links_analysis(self, doc, links_analyzer):
        """
        Add links analysis to the document.
        
        Args:
            doc (Document): The document
            links_analyzer (LinksAnalyzer): The links analyzer
        """
        doc.add_heading("Links Analysis", level=1)
        
        # Get the data
        links_df = links_analyzer.to_dataframe()
        
        # Add summary
        internal_links = len(links_df[links_df["Is Internal"]])
        external_links = len(links_df[links_df["Is External"]])
        broken_links = len(links_df[links_df["Is Broken"]])
        nofollow_links = len(links_df[links_df["Is Nofollow"]])
        
        doc.add_paragraph(f"Internal links: {internal_links}")
        doc.add_paragraph(f"External links: {external_links}")
        doc.add_paragraph(f"Broken links: {broken_links}")
        doc.add_paragraph(f"Nofollow links: {nofollow_links}")
        
        # Add broken links
        if broken_links > 0:
            doc.add_heading("Broken Links", level=2)
            
            broken_links_df = links_df[links_df["Is Broken"]]
            
            for _, row in broken_links_df.iterrows():
                doc.add_paragraph(f"{row['URL']}: {row['Link']}", style="List Bullet")
    
    def _add_reachability_analysis(self, doc, reachability_analyzer):
        """
        Add reachability analysis to the document.
        
        Args:
            doc (Document): The document
            reachability_analyzer (ReachabilityAnalyzer): The reachability analyzer
        """
        doc.add_heading("Reachability Analysis", level=1)
        
        # Get the data
        reachability_df = reachability_analyzer.to_dataframe()
        
        # Add summary
        reachable_urls = len(reachability_df[reachability_df["Is Reachable"]])
        orphan_pages = len(reachability_df[reachability_df["Is Orphan Page"]])
        
        doc.add_paragraph(f"Reachable URLs: {reachable_urls}")
        doc.add_paragraph(f"Orphan pages: {orphan_pages}")
        
        # Add orphan pages
        if orphan_pages > 0:
            doc.add_heading("Orphan Pages", level=2)
            
            orphan_pages_df = reachability_df[reachability_df["Is Orphan Page"]]
            
            for _, row in orphan_pages_df.iterrows():
                doc.add_paragraph(row["URL"], style="List Bullet")
        
        # Add clicks from home
        doc.add_heading("Clicks from Home", level=2)
        
        # Sort by clicks from home
        reachability_df = reachability_df.sort_values("Clicks from Home")
        
        # Add a table
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "URL"
        header_cells[1].text = "Clicks from Home"
        
        # Add data rows
        for _, row in reachability_df.iterrows():
            # Skip orphan pages
            if row["Is Orphan Page"]:
                continue
            
            # Add a row
            row_cells = table.add_row().cells
            row_cells[0].text = row["URL"]
            row_cells[1].text = str(row["Clicks from Home"])
    
    def _add_topic_analysis(self, doc, topic_analyzer):
        """
        Add topic analysis to the document.
        
        Args:
            doc (Document): The document
            topic_analyzer (MainTopicAnalyzer): The main topic analyzer
        """
        doc.add_heading("Topic Analysis", level=1)
        
        # Get the data
        topic_df = topic_analyzer.to_dataframe()
        
        # Add a word cloud of topics
        doc.add_heading("Topic Word Cloud", level=2)
        
        # Create a word cloud
        topics = " ".join(topic_df["Main Topics"].tolist())
        
        if topics:
            wordcloud = WordCloud(width=800, height=400, background_color="white").generate(topics)
            
            # Save the word cloud to a BytesIO object
            img_bytes = io.BytesIO()
            wordcloud.to_image().save(img_bytes, format="PNG")
            img_bytes.seek(0)
            
            # Add the word cloud to the document
            doc.add_picture(img_bytes, width=Inches(6))
        else:
            doc.add_paragraph("No topics found.")
        
        # Add a word cloud of keywords
        doc.add_heading("Keyword Word Cloud", level=2)
        
        # Create a word cloud
        keywords = " ".join(topic_df["Keywords"].tolist())
        
        if keywords:
            wordcloud = WordCloud(width=800, height=400, background_color="white").generate(keywords)
            
            # Save the word cloud to a BytesIO object
            img_bytes = io.BytesIO()
            wordcloud.to_image().save(img_bytes, format="PNG")
            img_bytes.seek(0)
            
            # Add the word cloud to the document
            doc.add_picture(img_bytes, width=Inches(6))
        else:
            doc.add_paragraph("No keywords found.")
        
        # Add topics per URL
        doc.add_heading("Topics per URL", level=2)
        
        # Add a table
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "URL"
        header_cells[1].text = "Main Topics"
        
        # Add data rows
        for _, row in topic_df.iterrows():
            # Add a row
            row_cells = table.add_row().cells
            row_cells[0].text = row["URL"]
            row_cells[1].text = row["Main Topics"]
    
    def _add_clusters_analysis(self, doc, clusters, item_type):
        """
        Add clusters analysis to the document.
        
        Args:
            doc (Document): The document
            clusters (list): The clusters
            item_type (str): The type of items in the clusters (Query or Keyword)
        """
        doc.add_heading(f"{item_type} Clusters", level=2)
        
        # Add a table
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Cluster"
        header_cells[1].text = f"Top {item_type}s"
        header_cells[2].text = "Metrics"
        
        # Add data rows
        for i, cluster in enumerate(clusters):
            # Add a row
            row_cells = table.add_row().cells
            row_cells[0].text = f"Cluster {i+1}"
            row_cells[1].text = ", ".join(cluster["top_queries" if item_type == "Query" else "top_keywords"])
            
            if item_type == "Query":
                row_cells[2].text = f"Queries: {cluster['queries']}\nImpressions: {cluster['impressions']:.0f}\nClicks: {cluster['clicks']:.0f}\nAvg Position: {cluster['avg_position']:.1f}"
            else:
                row_cells[2].text = f"Keywords: {cluster['keywords']}\nTraffic: {cluster['traffic']:.0f}\nAvg Position: {cluster['avg_position']:.1f}\nSearch Volume: {cluster['search_volume']:.0f}"
    
    def _add_topics_analysis(self, doc, topics):
        """
        Add topics analysis to the document.
        
        Args:
            doc (Document): The document
            topics (dict): The topics
        """
        doc.add_heading("Topics per URL", level=2)
        
        # Add a table
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "URL"
        header_cells[1].text = "Topics"
        
        # Add data rows
        for url, topic in topics.items():
            # Add a row
            row_cells = table.add_row().cells
            row_cells[0].text = url
            row_cells[1].text = ", ".join(topic)
    
    def _add_internal_link_suggestions(self, doc, suggestions, source_name=""):
        """
        Add internal link suggestions to the document.
        
        Args:
            doc (Document): The document
            suggestions (list): The internal link suggestions
            source_name (str, optional): The name of the source (e.g., "Search Console", "SEMrush")
        """
        if source_name:
            doc.add_heading(f"Internal Link Suggestions ({source_name})", level=2)
        else:
            doc.add_heading("Internal Link Suggestions", level=2)
        
        if not suggestions:
            doc.add_paragraph("No internal link suggestions found.")
            return
        
        # Add a table
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Source Page"
        header_cells[1].text = "Target Page"
        header_cells[2].text = "Topic"
        
        # Add data rows
        for suggestion in suggestions:
            # Add a row
            row_cells = table.add_row().cells
            row_cells[0].text = suggestion["source"]
            row_cells[1].text = suggestion["target"]
            row_cells[2].text = suggestion["topic"]
    
    def _add_visibility_traffic_analysis(self, doc, visibility, traffic):
        """
        Add visibility and traffic analysis to the document.
        
        Args:
            doc (Document): The document
            visibility (dict): The visibility per topic
            traffic (dict): The traffic per topic
        """
        doc.add_heading("Visibility and Traffic per Topic", level=2)
        
        # Add a table
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Topic"
        header_cells[1].text = "Visibility"
        header_cells[2].text = "Traffic"
        
        # Add data rows
        for topic in set(list(visibility.keys()) + list(traffic.keys())):
            # Add a row
            row_cells = table.add_row().cells
            row_cells[0].text = topic
            row_cells[1].text = str(visibility.get(topic, 0))
            row_cells[2].text = str(traffic.get(topic, 0))
    
    def _add_recommendations(self, doc, meta_analyzer, image_analyzer, links_analyzer, reachability_analyzer, topic_analyzer, search_console_analyzer=None, semrush_analyzer=None):
        """
        Add recommendations to the document.
        
        Args:
            doc (Document): The document
            meta_analyzer (MetaDescriptionAnalyzer): The meta description analyzer
            image_analyzer (ImageAnalyzer): The image analyzer
            links_analyzer (LinksAnalyzer): The links analyzer
            reachability_analyzer (ReachabilityAnalyzer): The reachability analyzer
            topic_analyzer (MainTopicAnalyzer): The main topic analyzer
            search_console_analyzer (SearchConsoleAnalyzer, optional): The Search Console analyzer
            semrush_analyzer (SEMrushAnalyzer, optional): The SEMrush analyzer
        """
        doc.add_heading("Recommendations", level=1)
        
        # Add meta description recommendations
        doc.add_heading("Meta Description Recommendations", level=2)
        
        meta_desc_df = meta_analyzer.to_dataframe()
        urls_without_meta_desc = len(meta_desc_df[~meta_desc_df["Has Meta Description"]])
        urls_with_short_meta_desc = len(meta_desc_df[meta_desc_df["Meta Description Quality"] == "too_short"])
        urls_with_long_meta_desc = len(meta_desc_df[meta_desc_df["Meta Description Quality"] == "too_long"])
        
        if urls_without_meta_desc > 0:
            doc.add_paragraph("Add meta descriptions to all pages without them.", style="List Bullet")
        
        if urls_with_short_meta_desc > 0:
            doc.add_paragraph("Improve short meta descriptions to be more descriptive.", style="List Bullet")
        
        if urls_with_long_meta_desc > 0:
            doc.add_paragraph("Shorten long meta descriptions to be under 160 characters.", style="List Bullet")
        
        # Add image recommendations
        doc.add_heading("Image Recommendations", level=2)
        
        image_df = image_analyzer.to_dataframe()
        images_without_alt = len(image_df[~image_df["Has Alt Text"]])
        large_images = len(image_df[image_df["Image Size"] > 100000])  # > 100 KB
        
        if images_without_alt > 0:
            doc.add_paragraph("Add alt text to all images without them.", style="List Bullet")
        
        if large_images > 0:
            doc.add_paragraph("Optimize large images to improve page load speed.", style="List Bullet")
        
        # Add link recommendations
        doc.add_heading("Link Recommendations", level=2)
        
        links_df = links_analyzer.to_dataframe()
        broken_links = len(links_df[links_df["Is Broken"]])
        
        if broken_links > 0:
            doc.add_paragraph("Fix all broken links.", style="List Bullet")
        
        # Add reachability recommendations
        doc.add_heading("Reachability Recommendations", level=2)
        
        reachability_df = reachability_analyzer.to_dataframe()
        orphan_pages = len(reachability_df[reachability_df["Is Orphan Page"]])
        
        if orphan_pages > 0:
            doc.add_paragraph("Add internal links to orphan pages to make them reachable from the home page.", style="List Bullet")
        
        # Add topic recommendations
        doc.add_heading("Topic Recommendations", level=2)
        
        doc.add_paragraph("Ensure that each page focuses on a specific topic to improve relevance.", style="List Bullet")
        doc.add_paragraph("Use relevant keywords in headings, content, and meta descriptions.", style="List Bullet")
        
        # Add Search Console recommendations if available
        if search_console_analyzer:
            doc.add_heading("Search Console Recommendations", level=2)
            
            doc.add_paragraph("Focus on improving rankings for queries with high impressions but low clicks.", style="List Bullet")
            doc.add_paragraph("Create content around related queries to expand your topical coverage.", style="List Bullet")
        
        # Add SEMrush recommendations if available
        if semrush_analyzer:
            doc.add_heading("SEMrush Recommendations", level=2)
            
            doc.add_paragraph("Target keywords with high search volume and low competition.", style="List Bullet")
            doc.add_paragraph("Improve content for keywords with low visibility but high traffic potential.", style="List Bullet")
        
        # Add internal link recommendations
        doc.add_heading("Internal Linking Recommendations", level=2)
        
        doc.add_paragraph("Implement a logical site structure with clear navigation.", style="List Bullet")
        doc.add_paragraph("Use descriptive anchor text for internal links.", style="List Bullet")
        doc.add_paragraph("Link related content together to improve topical relevance.", style="List Bullet")
        
        if search_console_analyzer or semrush_analyzer:
            doc.add_paragraph("Implement the internal link suggestions provided in this report.", style="List Bullet")
    
    def _add_query_comparison(self, doc, query_comparison):
        """
        Add query comparison to the document.
        
        Args:
            doc (Document): The document
            query_comparison (dict): The query comparison
        """
        doc.add_heading("Query Comparison", level=2)
        
        if not query_comparison:
            doc.add_paragraph("No query comparison data available.")
            return
        
        # Add a table
        table = doc.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Query"
        header_cells[1].text = "Impressions Before"
        header_cells[2].text = "Impressions After"
        header_cells[3].text = "Clicks Before"
        header_cells[4].text = "Clicks After"
        header_cells[5].text = "Position Before"
        header_cells[6].text = "Position After"
        
        # Add data rows
        for query, data in query_comparison.items():
            # Add a row
            row_cells = table.add_row().cells
            row_cells[0].text = query
            row_cells[1].text = str(data.get("impressions_before", 0))
            row_cells[2].text = str(data.get("impressions_after", 0))
            row_cells[3].text = str(data.get("clicks_before", 0))
            row_cells[4].text = str(data.get("clicks_after", 0))
            row_cells[5].text = f"{data.get('position_before', 0):.1f}"
            row_cells[6].text = f"{data.get('position_after', 0):.1f}"
    
    def _add_landing_page_comparison(self, doc, landing_page_comparison):
        """
        Add landing page comparison to the document.
        
        Args:
            doc (Document): The document
            landing_page_comparison (dict): The landing page comparison
        """
        doc.add_heading("Landing Page Comparison", level=2)
        
        if not landing_page_comparison:
            doc.add_paragraph("No landing page comparison data available.")
            return
        
        # Add a table
        table = doc.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Landing Page"
        header_cells[1].text = "Impressions Before"
        header_cells[2].text = "Impressions After"
        header_cells[3].text = "Clicks Before"
        header_cells[4].text = "Clicks After"
        header_cells[5].text = "Position Before"
        header_cells[6].text = "Position After"
        
        # Add data rows
        for landing_page, data in landing_page_comparison.items():
            # Add a row
            row_cells = table.add_row().cells
            row_cells[0].text = landing_page
            row_cells[1].text = str(data.get("impressions_before", 0))
            row_cells[2].text = str(data.get("impressions_after", 0))
            row_cells[3].text = str(data.get("clicks_before", 0))
            row_cells[4].text = str(data.get("clicks_after", 0))
            row_cells[5].text = f"{data.get('position_before', 0):.1f}"
            row_cells[6].text = f"{data.get('position_after', 0):.1f}"
    
    def _add_improved_queries(self, doc, improved_queries):
        """
        Add improved queries to the document.
        
        Args:
            doc (Document): The document
            improved_queries (list): The improved queries
        """
        doc.add_heading("Improved Queries", level=2)
        
        if not improved_queries:
            doc.add_paragraph("No improved queries data available.")
            return
        
        # Add a table
        table = doc.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Query"
        header_cells[1].text = "Impressions Before"
        header_cells[2].text = "Impressions After"
        header_cells[3].text = "Clicks Before"
        header_cells[4].text = "Clicks After"
        header_cells[5].text = "Position Before"
        header_cells[6].text = "Position After"
        
        # Add data rows
        for data in improved_queries:
            # Add a row
            row_cells = table.add_row().cells
            row_cells[0].text = data.get("query", "")
            row_cells[1].text = str(data.get("impressions_before", 0))
            row_cells[2].text = str(data.get("impressions_after", 0))
            row_cells[3].text = str(data.get("clicks_before", 0))
            row_cells[4].text = str(data.get("clicks_after", 0))
            row_cells[5].text = f"{data.get('position_before', 0):.1f}"
            row_cells[6].text = f"{data.get('position_after', 0):.1f}"
    
    def _add_declined_queries(self, doc, declined_queries):
        """
        Add declined queries to the document.
        
        Args:
            doc (Document): The document
            declined_queries (list): The declined queries
        """
        doc.add_heading("Declined Queries", level=2)
        
        if not declined_queries:
            doc.add_paragraph("No declined queries data available.")
            return
        
        # Add a table
        table = doc.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Query"
        header_cells[1].text = "Impressions Before"
        header_cells[2].text = "Impressions After"
        header_cells[3].text = "Clicks Before"
        header_cells[4].text = "Clicks After"
        header_cells[5].text = "Position Before"
        header_cells[6].text = "Position After"
        
        # Add data rows
        for data in declined_queries:
            # Add a row
            row_cells = table.add_row().cells
            row_cells[0].text = data.get("query", "")
            row_cells[1].text = str(data.get("impressions_before", 0))
            row_cells[2].text = str(data.get("impressions_after", 0))
            row_cells[3].text = str(data.get("clicks_before", 0))
            row_cells[4].text = str(data.get("clicks_after", 0))
            row_cells[5].text = f"{data.get('position_before', 0):.1f}"
            row_cells[6].text = f"{data.get('position_after', 0):.1f}"
    
    def _add_improved_landing_pages(self, doc, improved_landing_pages):
        """
        Add improved landing pages to the document.
        
        Args:
            doc (Document): The document
            improved_landing_pages (list): The improved landing pages
        """
        doc.add_heading("Improved Landing Pages", level=2)
        
        if not improved_landing_pages:
            doc.add_paragraph("No improved landing pages data available.")
            return
        
        # Add a table
        table = doc.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Landing Page"
        header_cells[1].text = "Impressions Before"
        header_cells[2].text = "Impressions After"
        header_cells[3].text = "Clicks Before"
        header_cells[4].text = "Clicks After"
        header_cells[5].text = "Position Before"
        header_cells[6].text = "Position After"
        
        # Add data rows
        for data in improved_landing_pages:
            # Add a row
            row_cells = table.add_row().cells
            row_cells[0].text = data.get("landing_page", "")
            row_cells[1].text = str(data.get("impressions_before", 0))
            row_cells[2].text = str(data.get("impressions_after", 0))
            row_cells[3].text = str(data.get("clicks_before", 0))
            row_cells[4].text = str(data.get("clicks_after", 0))
            row_cells[5].text = f"{data.get('position_before', 0):.1f}"
            row_cells[6].text = f"{data.get('position_after', 0):.1f}"
    
    def _add_declined_landing_pages(self, doc, declined_landing_pages):
        """
        Add declined landing pages to the document.
        
        Args:
            doc (Document): The document
            declined_landing_pages (list): The declined landing pages
        """
        doc.add_heading("Declined Landing Pages", level=2)
        
        if not declined_landing_pages:
            doc.add_paragraph("No declined landing pages data available.")
            return
        
        # Add a table
        table = doc.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Landing Page"
        header_cells[1].text = "Impressions Before"
        header_cells[2].text = "Impressions After"
        header_cells[3].text = "Clicks Before"
        header_cells[4].text = "Clicks After"
        header_cells[5].text = "Position Before"
        header_cells[6].text = "Position After"
        
        # Add data rows
        for data in declined_landing_pages:
            # Add a row
            row_cells = table.add_row().cells
            row_cells[0].text = data.get("landing_page", "")
            row_cells[1].text = str(data.get("impressions_before", 0))
            row_cells[2].text = str(data.get("impressions_after", 0))
            row_cells[3].text = str(data.get("clicks_before", 0))
            row_cells[4].text = str(data.get("clicks_after", 0))
            row_cells[5].text = f"{data.get('position_before', 0):.1f}"
            row_cells[6].text = f"{data.get('position_after', 0):.1f}"
    
    def _add_executive_summary(self, doc, url, meta_analyzer, image_analyzer, links_analyzer, reachability_analyzer, topic_analyzer, search_console_analyzer=None, semrush_analyzer=None):
        """
        Add an executive summary to the document.
        
        Args:
            doc (Document): The document
            url (str): The URL
            meta_analyzer (MetaDescriptionAnalyzer): The meta description analyzer
            image_analyzer (ImageAnalyzer): The image analyzer
            links_analyzer (LinksAnalyzer): The links analyzer
            reachability_analyzer (ReachabilityAnalyzer): The reachability analyzer
            topic_analyzer (MainTopicAnalyzer): The main topic analyzer
            search_console_analyzer (SearchConsoleAnalyzer, optional): The Search Console analyzer
            semrush_analyzer (SEMrushAnalyzer, optional): The SEMrush analyzer
        """
        doc.add_heading("Executive Summary", level=1)
        
        # Add URL
        doc.add_paragraph(f"URL: {url}")
        
        # Add summary paragraph
        doc.add_paragraph("This report provides a comprehensive analysis of the SEO performance of the website. It includes an analysis of meta descriptions, images, links, reachability, and topics. The report also provides recommendations for improving the SEO performance of the website.")
        
        # Add key findings
        doc.add_heading("Key Findings", level=2)
        
        # Meta description findings
        meta_desc_df = meta_analyzer.to_dataframe()
        urls_with_meta_desc = len(meta_desc_df[meta_desc_df["Has Meta Description"]])
        urls_without_meta_desc = len(meta_desc_df) - urls_with_meta_desc
        
        if urls_without_meta_desc > 0:
            doc.add_paragraph(f"{urls_without_meta_desc} pages don't have meta descriptions.", style="List Bullet")
        
        # Image findings
        image_df = image_analyzer.to_dataframe()
        images_without_alt = len(image_df[~image_df["Has Alt Text"]])
        
        if images_without_alt > 0:
            doc.add_paragraph(f"{images_without_alt} images don't have alt text.", style="List Bullet")
        
        # Link findings
        links_df = links_analyzer.to_dataframe()
        broken_links = len(links_df[links_df["Is Broken"]])
        
        if broken_links > 0:
            doc.add_paragraph(f"{broken_links} links are broken.", style="List Bullet")
        
        # Reachability findings
        reachability_df = reachability_analyzer.to_dataframe()
        orphan_pages = len(reachability_df[reachability_df["Is Orphan Page"]])
        
        if orphan_pages > 0:
            doc.add_paragraph(f"{orphan_pages} pages are orphan pages.", style="List Bullet")
        
        # Search Console findings
        if search_console_analyzer:
            doc.add_paragraph("Search Console data has been analyzed to identify opportunities for improvement.", style="List Bullet")
        
        # SEMrush findings
        if semrush_analyzer:
            doc.add_paragraph("SEMrush data has been analyzed to identify keyword opportunities.", style="List Bullet")
        
        # Add key recommendations
        doc.add_heading("Key Recommendations", level=2)
        
        if urls_without_meta_desc > 0:
            doc.add_paragraph("Add meta descriptions to all pages without them.", style="List Bullet")
        
        if images_without_alt > 0:
            doc.add_paragraph("Add alt text to all images without them.", style="List Bullet")
        
        if broken_links > 0:
            doc.add_paragraph("Fix all broken links.", style="List Bullet")
        
        if orphan_pages > 0:
            doc.add_paragraph("Add internal links to orphan pages to make them reachable from the home page.", style="List Bullet")
        
        doc.add_paragraph("Implement a logical site structure with clear navigation.", style="List Bullet")
        doc.add_paragraph("Use descriptive anchor text for internal links.", style="List Bullet")
        doc.add_paragraph("Ensure that each page focuses on a specific topic to improve relevance.", style="List Bullet")
